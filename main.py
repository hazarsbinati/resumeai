import os
from docx import Document
from transformers import pipeline

def get_docx_files(directory):
    return [os.path.join(directory, f) for f in os.listdir(directory) if f.lower().endswith('.docx')]

def extract_sections(doc):
    sections = {}
    current_heading = None
    current_content = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if current_heading and current_content:
                sections.setdefault(current_heading, []).append('\n'.join(current_content).strip())
            current_heading = para.text.strip()
            current_content = []
        else:
            current_content.append(para.text.strip())
    if current_heading and current_content:
        sections.setdefault(current_heading, []).append('\n'.join(current_content).strip())
    return sections

def merge_and_summarize_sections(all_sections):
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    final_sections = {}
    for heading, contents in all_sections.items():
        unique_contents = list(set(contents))
        if len(unique_contents) == 1:
            final_sections[heading] = unique_contents[0]
        else:
            # Summarize repeated sections
            combined = "\n".join(unique_contents)
            # transformers pipeline has a max_token limit
            if len(combined) > 2000:
                pieces = [combined[i:i+2000] for i in range(0, len(combined), 2000)]
                summarized = " ".join([summarizer(piece, max_length=150, min_length=40, do_sample=False)[0]['summary_text'] for piece in pieces])
            else:
                summarized = summarizer(combined, max_length=150, min_length=40, do_sample=False)[0]['summary_text']
            final_sections[heading] = summarized
    return final_sections

def write_docx(sections, output_file):
    doc = Document()
    for heading, content in sections.items():
        doc.add_heading(heading, level=1)
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    doc.save(output_file)

if __name__ == "__main__":
    directory = "/Users/hazarsbinati/icloud drive/resume"  # Change as needed
    output_file = "final_resume.docx"
    all_sections = {}
    for file in get_docx_files(directory):
        doc = Document(file)
        sections = extract_sections(doc)
        for heading, content_list in sections.items():
            all_sections.setdefault(heading, []).extend(content_list)
    merged_sections = merge_and_summarize_sections(all_sections)
    write_docx(merged_sections, output_file)
    print(f"Final resume written to {output_file}")
